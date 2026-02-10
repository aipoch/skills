#!/usr/bin/env python3
"""
Blind Review Sanitizer (ID: 162)
一键移除文稿中的作者姓名、单位、致谢及过多的自引，符合双盲评审要求。
"""

import argparse
import re
import sys
from pathlib import Path
from typing import List, Optional, Tuple


# 常见机构关键词
INSTITUTION_KEYWORDS = [
    r'University', r'College', r'Institute', r'Academy', r'School',
    r'大学', r'学院', r'研究所', r'研究院', r'实验室', r'Lab\.?',
    r'Department', r'Dept\.?', r'系', r'中心', r'Center', r'Centre'
]

# 致谢相关标题
ACKNOWLEDGMENT_TITLES = [
    r'(?i)^\s*acknowledgments?\s*$',
    r'(?i)^\s*致谢\s*$',
    r'(?i)^\s*acknowledgements?\s*$',
    r'(?i)^\s*funding\s*$',
    r'(?i)^\s*资助\s*$',
]

# 自引检测模式
SELF_CITATION_PATTERNS = [
    r'\bour\s+(?:previous|prior|earlier)\s+(?:work|study|research|paper)s?\b',
    r'\bwe\s+(?:previously|earlier)\s+(?:showed|demonstrated|reported|found)\b',
    r'\bin\s+our\s+(?:previous|prior)\s+(?:work|study|research)\b',
    r'\b(?:my|our)\s+(?:own\s+)?(?:previous|prior|earlier)\b',
]

# 邮箱模式
EMAIL_PATTERN = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'

# 电话模式
PHONE_PATTERN = r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}\b'


class BlindReviewSanitizer:
    """双盲评审脱敏处理器"""
    
    def __init__(
        self,
        authors: Optional[List[str]] = None,
        keep_acknowledgments: bool = False,
        highlight_self_cites: bool = False
    ):
        self.authors = [a.strip() for a in (authors or []) if a.strip()]
        self.keep_acknowledgments = keep_acknowledgments
        self.highlight_self_cites = highlight_self_cites
        self.removed_items = []
        
    def sanitize_text(self, text: str) -> str:
        """对纯文本进行脱敏处理"""
        result = text
        
        # 1. 移除作者姓名
        result = self._remove_author_names(result)
        
        # 2. 移除单位信息
        result = self._remove_institutions(result)
        
        # 3. 移除联系信息
        result = self._remove_contact_info(result)
        
        # 4. 处理自引
        result = self._handle_self_citations(result)
        
        return result
    
    def _remove_author_names(self, text: str) -> str:
        """移除作者姓名"""
        if not self.authors:
            return text
            
        result = text
        for author in self.authors:
            # 匹配全名和姓氏
            patterns = [
                re.escape(author),
            ]
            for pattern in patterns:
                result = re.sub(
                    pattern,
                    '[AUTHOR NAME]',
                    result,
                    flags=re.IGNORECASE
                )
        return result
    
    def _remove_institutions(self, text: str) -> str:
        """移除机构/单位信息"""
        result = text
        
        # 匹配常见机构格式
        for keyword in INSTITUTION_KEYWORDS:
            # 匹配 "XX University", "University of XX", "XX学院" 等
            pattern = rf'\b[A-Z][A-Za-z\s]*{keyword}[A-Za-z\s]*\b|\b{keyword}[\u4e00-\u9fa5]+\b'
            matches = re.finditer(pattern, result, re.IGNORECASE)
            for match in list(matches)[::-1]:  # 反向替换避免位置变化
                self.removed_items.append(f"Institution: {match.group()}")
                result = result[:match.start()] + '[INSTITUTION]' + result[match.end():]
        
        return result
    
    def _remove_contact_info(self, text: str) -> str:
        """移除联系信息（邮箱、电话）"""
        result = text
        
        # 移除邮箱
        matches = re.finditer(EMAIL_PATTERN, result, re.IGNORECASE)
        for match in list(matches)[::-1]:
            self.removed_items.append(f"Email: {match.group()}")
            result = result[:match.start()] + '[EMAIL]' + result[match.end():]
        
        # 移除电话
        matches = re.finditer(PHONE_PATTERN, result)
        for match in list(matches)[::-1]:
            self.removed_items.append(f"Phone: {match.group()}")
            result = result[:match.start()] + '[PHONE]' + result[match.end():]
        
        return result
    
    def _handle_self_citations(self, text: str) -> str:
        """处理自引"""
        result = text
        
        for pattern in SELF_CITATION_PATTERNS:
            if self.highlight_self_cites:
                result = re.sub(pattern, r'[SELF-CITE: \g<0>]', result, flags=re.IGNORECASE)
            else:
                def replace_func(match):
                    self.removed_items.append(f"Self-citation: {match.group()}")
                    return '[PREVIOUS WORK]'
                result = re.sub(pattern, replace_func, result, flags=re.IGNORECASE)
        
        return result
    
    def remove_acknowledgments(self, lines: List[str]) -> List[str]:
        """移除致谢部分"""
        if self.keep_acknowledgments:
            return lines
            
        result = []
        in_acknowledgment = False
        
        for line in lines:
            # 检测致谢标题
            if any(re.match(pattern, line.strip()) for pattern in ACKNOWLEDGMENT_TITLES):
                in_acknowledgment = True
                self.removed_items.append(f"Acknowledgment section removed")
                result.append('[ACKNOWLEDGMENTS REMOVED]\n')
                continue
            
            # 如果在致谢部分，检测是否结束（遇到下一个标题）
            if in_acknowledgment:
                if re.match(r'^[#\d\s]', line.strip()) or line.strip().isupper():
                    in_acknowledgment = False
                else:
                    continue
            
            result.append(line)
        
        return result


class DocxProcessor:
    """DOCX文档处理器"""
    
    def __init__(self, sanitizer: BlindReviewSanitizer):
        self.sanitizer = sanitizer
    
    def process(self, input_path: Path, output_path: Path) -> None:
        """处理DOCX文件"""
        try:
            from docx import Document
        except ImportError:
            print("Error: python-docx not installed. Run: pip install python-docx")
            sys.exit(1)
        
        doc = Document(input_path)
        
        # 处理段落
        for para in doc.paragraphs:
            # 检测并移除致谢
            if not self.sanitizer.keep_acknowledgments:
                if any(re.match(pattern, para.text.strip()) for pattern in ACKNOWLEDGMENT_TITLES):
                    para.text = '[ACKNOWLEDGMENTS REMOVED]'
                    continue
            
            # 处理文本内容
            if para.text.strip():
                para.text = self.sanitizer.sanitize_text(para.text)
        
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        cell.text = self.sanitizer.sanitize_text(cell.text)
        
        # 保存
        doc.save(output_path)


class TxtProcessor:
    """纯文本处理器"""
    
    def __init__(self, sanitizer: BlindReviewSanitizer):
        self.sanitizer = sanitizer
    
    def process(self, input_path: Path, output_path: Path) -> None:
        """处理文本文件"""
        with open(input_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # 移除致谢
        lines = self.sanitizer.remove_acknowledgments(lines)
        
        # 处理每行内容
        result_lines = []
        for line in lines:
            processed = self.sanitizer.sanitize_text(line)
            result_lines.append(processed)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.writelines(result_lines)


class MdProcessor(TxtProcessor):
    """Markdown处理器（继承TxtProcessor，保留特殊处理扩展性）"""
    pass


def get_processor(file_path: Path, sanitizer: BlindReviewSanitizer):
    """根据文件类型获取相应处理器"""
    suffix = file_path.suffix.lower()
    
    if suffix == '.docx':
        return DocxProcessor(sanitizer)
    elif suffix == '.md':
        return MdProcessor(sanitizer)
    elif suffix == '.txt':
        return TxtProcessor(sanitizer)
    else:
        raise ValueError(f"Unsupported file format: {suffix}")


def main():
    parser = argparse.ArgumentParser(
        description='Blind Review Sanitizer - 双盲评审文稿脱敏工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  %(prog)s --input paper.docx
  %(prog)s --input paper.md --authors "张三,李四" --output blinded.md
  %(prog)s --input paper.txt --keep-acknowledgments --highlight-self-cites
        """
    )
    
    parser.add_argument('--input', '-i', required=True, help='输入文件路径 (docx/txt/md)')
    parser.add_argument('--output', '-o', help='输出文件路径（默认添加 -blinded 后缀）')
    parser.add_argument('--authors', help='作者姓名，逗号分隔')
    parser.add_argument('--keep-acknowledgments', action='store_true', help='保留致谢部分')
    parser.add_argument('--highlight-self-cites', action='store_true', help='仅高亮自引而不替换')
    
    args = parser.parse_args()
    
    # 解析输入路径
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)
    
    # 确定输出路径
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.parent / f"{input_path.stem}-blinded{input_path.suffix}"
    
    # 解析作者列表
    authors = None
    if args.authors:
        authors = [a.strip() for a in args.authors.split(',') if a.strip()]
    
    # 创建脱敏器
    sanitizer = BlindReviewSanitizer(
        authors=authors,
        keep_acknowledgments=args.keep_acknowledgments,
        highlight_self_cites=args.highlight_self_cites
    )
    
    # 获取处理器
    try:
        processor = get_processor(input_path, sanitizer)
    except ValueError as e:
        print(f"Error: {e}")
        sys.exit(1)
    
    # 处理文件
    print(f"Processing: {input_path}")
    processor.process(input_path, output_path)
    
    # 输出统计
    print(f"Output: {output_path}")
    print(f"Items processed: {len(sanitizer.removed_items)}")
    if sanitizer.removed_items:
        print("Summary:")
        for item in set(sanitizer.removed_items):
            count = sanitizer.removed_items.count(item)
            if count > 1:
                print(f"  - {item} (x{count})")
            else:
                print(f"  - {item}")
    
    print("Done!")


if __name__ == '__main__':
    main()
